from pathlib import Path

from docx import Document as DOCXDocument

from src.schemas import Document, Page


class DOCXExtractor:

    def can_process(self, document: Path) -> bool:
        return document.suffix.lower() == ".docx"

    def extract(self, document: Path) -> Document:
        pages: list[Page] = []

        docx_document = DOCXDocument(document)

        for paragraph_number, paragraph in enumerate(
            docx_document.paragraphs,
            start=1,
        ):
            pages.append(
                Page(
                    page_number=paragraph_number,
                    raw_text=paragraph.text.strip(),
                    extraction_method="docx",
                )
            )

        return Document(
            document_id=document.stem,
            source_path=str(document),
            file_type="docx",
            pages=pages,
            metadata={},
        )


if __name__ == "__main__":
    from pathlib import Path

    from src.config import DOCX_TEST_FILE

    docx_extractor = DOCXExtractor()
    document_path = Path(DOCX_TEST_FILE)

    if docx_extractor.can_process(document_path):
        result = docx_extractor.extract(document_path)

        for page in result.pages:
            print(f"\n{'=' * 60}")
            print(f"SOURCE {page.page_number}")
            print(f"{'=' * 60}")
            print(page.raw_text)
    else:
        print("File is not .docx")